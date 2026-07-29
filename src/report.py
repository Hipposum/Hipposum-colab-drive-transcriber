"""
report.py — сборка читаемого транскрипта в стиле SpeechToText:

    19.07.26_08.48.mp4

    Спикер 1:
    00:00:21 - Текст реплики, склеенный в абзац...
    00:01:02 - Следующий абзац того же спикера...

    Спикер 2:
    00:01:52 - Нет.

Реплики одного спикера группируются в абзацы: новый абзац начинается после
заметной паузы либо когда абзац становится слишком длинным. Таймкод — начало абзаца.
Дополнительно умеет собирать тот же транскрипт в .docx (python-docx).
"""

# ── Правила разбивки на абзацы (подобраны по эталонному файлу SpeechToText) ──
SPLIT_PAUSE_SEC     = 2.5   # пауза между репликами → новый абзац
MAX_PARAGRAPH_SEC   = 45.0  # длительность абзаца → принудительный разрыв
MAX_PARAGRAPH_CHARS = 450   # длина абзаца в символах → принудительный разрыв

UNKNOWN_LABEL = "Спикер ?"


def _fmt_hms(seconds):
    seconds = max(0, int(seconds))
    h, rest = divmod(seconds, 3600)
    m, s = divmod(rest, 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def _speaker_label(raw, mapping):
    """SPEAKER_00 → 'Спикер 1' (нумерация по порядку первого появления)."""
    if not raw or raw == "UNKNOWN":
        return UNKNOWN_LABEL
    if raw not in mapping:
        mapping[raw] = f"Спикер {len(mapping) + 1}"
    return mapping[raw]


def build_speaker_blocks(segments, cfg=None):
    """
    Группирует сегменты в блоки для рендера.
    Возвращает список dict: {speaker, paragraphs: [{start, text}]}.
    Блок = подряд идущие абзацы одного спикера; абзац = склейка реплик
    до паузы/лимита длины.
    """
    cfg = cfg or {}
    split_pause = cfg.get("report_split_pause_sec", SPLIT_PAUSE_SEC)
    max_sec     = cfg.get("report_max_paragraph_sec", MAX_PARAGRAPH_SEC)
    max_chars   = cfg.get("report_max_paragraph_chars", MAX_PARAGRAPH_CHARS)

    mapping = {}
    blocks = []
    cur_block = None   # {"speaker": label, "paragraphs": [...]}
    cur_par = None     # {"start": t, "end": t, "text": str}

    def close_par():
        nonlocal cur_par
        if cur_par and cur_par["text"].strip():
            cur_block["paragraphs"].append(
                {"start": cur_par["start"], "text": cur_par["text"].strip()})
        cur_par = None

    for seg in segments:
        text = (seg.get("text") or "").strip()
        if not text:
            continue
        start = seg.get("start", 0) or 0
        end = seg.get("end", start) or start
        label = (UNKNOWN_LABEL if seg.get("_is_placeholder")
                 else _speaker_label(seg.get("speaker"), mapping))

        if cur_block is None or cur_block["speaker"] != label:
            if cur_block is not None:
                close_par()
                blocks.append(cur_block)
            cur_block = {"speaker": label, "paragraphs": []}
            cur_par = None

        if cur_par is None:
            cur_par = {"start": start, "end": end, "text": text}
            continue

        gap = start - cur_par["end"]
        too_long = (end - cur_par["start"] > max_sec
                    or len(cur_par["text"]) + len(text) > max_chars)
        if gap > split_pause or too_long:
            close_par()
            cur_par = {"start": start, "end": end, "text": text}
        else:
            cur_par["text"] += " " + text
            cur_par["end"] = end

    if cur_block is not None:
        close_par()
        blocks.append(cur_block)
    return blocks


def build_full_document(analytics, llm_result, segments, audio_duration, file_name="", path=None):
    """Транскрипт в текстовом виде (главный .txt). analytics/llm_result не используются —
    сигнатура сохранена для совместимости со stages.py."""
    blocks = build_speaker_blocks(segments)
    lines = []
    if file_name:
        lines.append(file_name)
    for block in blocks:
        lines.append("")
        lines.append(f"{block['speaker']}:")
        for par in block["paragraphs"]:
            lines.append(f"{_fmt_hms(par['start'])} - {par['text']}")
    return "\n".join(lines) + "\n"


def build_docx(segments, file_name, out_path):
    """
    Тот же транскрипт в .docx (как отдаёт SpeechToText):
    имя файла справа курсивом, «Спикер N:» жирным, абзацы с таймкодами.
    Требует python-docx; при его отсутствии — ImportError (вызывающий код решает сам).
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(file_name)
    run.italic = True

    for block in build_speaker_blocks(segments):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        run = p.add_run(f"{block['speaker']}:")
        run.bold = True
        for par in block["paragraphs"]:
            doc.add_paragraph(f"{_fmt_hms(par['start'])} - {par['text']}")

    doc.save(out_path)
    return out_path
